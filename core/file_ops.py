"""Read PDF, docx, xlsx, and text files for the assistant."""
import os
from typing import Tuple
from core.logger import get_logger
log = get_logger("file_ops")
MAX_CHARS = 8000
TEXT_EXTS = {".txt",".md",".py",".js",".ts",".jsx",".tsx",".json",".yaml",".yml",
             ".toml",".csv",".log",".html",".css",".sql",".sh",".bat",".env",
             ".ini",".cfg",".xml",".rs",".go",".java",".c",".cpp",".h"}

def read_file(path: str) -> Tuple[str, str]:
    path = path.strip().strip('"\'')
    if not os.path.exists(path): return f"File not found: {path}", "error"
    ext = os.path.splitext(path)[1].lower()
    try:
        if ext in TEXT_EXTS:   return _text(path), "text"
        elif ext == ".pdf":    return _pdf(path), "pdf"
        elif ext == ".docx":   return _docx(path), "docx"
        elif ext in (".xlsx",".xls"): return _xlsx(path), "xlsx"
        else:                  return _text(path), "text"
    except Exception as e: return f"Read error: {e}", "error"

def _text(path):
    fname = os.path.basename(path); size = os.path.getsize(path)//1024
    with open(path, "r", encoding="utf-8", errors="replace") as f: content = f.read()
    hdr = f"File: {fname} ({size}KB)\n{'─'*40}\n"
    return hdr + (content[:MAX_CHARS] + "\n[truncated]" if len(content) > MAX_CHARS else content)

def _pdf(path):
    try: import pdfplumber
    except ImportError: return "pdfplumber not installed. Run: pip install pdfplumber"
    fname = os.path.basename(path); parts = []
    with pdfplumber.open(path) as pdf:
        total = len(pdf.pages)
        for i, page in enumerate(pdf.pages):
            text = page.extract_text() or ""
            if text.strip(): parts.append(f"[Page {i+1}]\n{text}")
            if sum(len(p) for p in parts) > MAX_CHARS:
                parts.append(f"[truncated at page {i+1} of {total}]"); break
    return f"PDF: {fname} ({total} pages)\n{'─'*40}\n" + "\n\n".join(parts)

def _docx(path):
    try: from docx import Document
    except ImportError: return "python-docx not installed. Run: pip install python-docx"
    doc = Document(path); fname = os.path.basename(path)
    content = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    return f"Word: {fname}\n{'─'*40}\n" + (content[:MAX_CHARS]+"[truncated]" if len(content)>MAX_CHARS else content)

def _xlsx(path):
    try: import openpyxl
    except ImportError: return "openpyxl not installed. Run: pip install openpyxl"
    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    fname = os.path.basename(path); parts = []
    for ws in wb.worksheets:
        rows = []
        for row in ws.iter_rows(max_row=200, values_only=True):
            if any(c is not None for c in row):
                rows.append("\t".join(str(c) if c is not None else "" for c in row))
        if rows: parts.append(f"Sheet: {ws.title}\n" + "\n".join(rows[:100]))
        if sum(len(p) for p in parts) > MAX_CHARS: break
    wb.close()
    return f"Excel: {fname}\n{'─'*40}\n" + "\n\n".join(parts)

# ── Write ────────────────────────────────────────────────────────────────────
_FORBIDDEN = ("c:\\windows", "c:\\program files", "c:\\programdata")


def _default_dir() -> str:
    home = os.path.expanduser("~")
    for d in (os.path.join(home, "Desktop"), os.path.join(home, "OneDrive", "Desktop")):
        if os.path.isdir(d):
            return d
    return home


def write_file(path: str, content: str, confirmed: bool = False) -> str:
    """Create or overwrite a file. Plain text for most extensions, .docx via
    python-docx. Relative paths land on the Desktop."""
    path = os.path.expanduser(path.strip().strip('"\''))
    if not os.path.isabs(path):
        path = os.path.join(_default_dir(), path)
    import config as cfg
    if not confirmed and any(path.lower().startswith(p) for p in _FORBIDDEN):
        if not cfg.get("full_control"):
            return (f"Refusing to write into a system directory: {path} "
                    "(enable 'Full computer control' in Settings to allow)")
        if cfg.get("confirm_destructive"):
            from core import confirm
            confirm.set_pending("write", (path, content), f"write into {path}")
            return (f"⚠️ CONFIRM NEEDED — writing into a system directory:\n{path}\n"
                    "Reply 'yes' to proceed or 'no' to cancel.")
    try:
        existed = os.path.exists(path)
        parent = os.path.dirname(path)
        if parent:
            os.makedirs(parent, exist_ok=True)
        ext = os.path.splitext(path)[1].lower()
        if ext == ".docx":
            try:
                from docx import Document
            except ImportError:
                return "python-docx not installed. Run: pip install python-docx"
            doc = Document()
            for line in content.split("\n"):
                doc.add_paragraph(line)
            doc.save(path)
        elif ext == ".pdf":
            return "Writing PDFs isn't supported — use .docx, .md or .txt instead."
        else:
            with open(path, "w", encoding="utf-8") as f:
                f.write(content)
        verb = "Overwrote" if existed else "Created"
        log.info(f"{verb} {path}")
        return f"✓ {verb} {path} ({len(content)} chars)"
    except Exception as e:
        return f"Write error: {e}"


# ── Find ─────────────────────────────────────────────────────────────────────
_SKIP_DIRS = {"appdata", "node_modules", ".git", "venv", ".venv", "__pycache__",
              "site-packages", "$recycle.bin", "windows", "program files",
              "programdata", "library", "cache", ".cache"}


def find_files(query: str, root: str = None, max_results: int = 20,
               time_budget: float = 10.0) -> str:
    """Find files by name (substring or glob) under the user's folders."""
    import fnmatch
    import time
    query = (query or "").strip().strip('"\'')
    if not query:
        return "FIND needs a file name or pattern, e.g. [FIND: report.docx] or [FIND: *.pdf]"
    pattern = query if any(ch in query for ch in "*?") else f"*{query}*"
    pattern = pattern.lower()

    home = os.path.expanduser("~")
    if root:
        roots = [os.path.expanduser(root.strip().strip('"\''))]
    else:
        roots, seen = [], set()
        for d in ("Desktop", "Documents", "Downloads", "Pictures", "Videos", "Music"):
            for base in (home, os.path.join(home, "OneDrive")):
                p = os.path.join(base, d)
                if os.path.isdir(p) and os.path.realpath(p) not in seen:
                    seen.add(os.path.realpath(p))
                    roots.append(p)

    deadline = time.time() + time_budget
    matches, timed_out = [], False
    for r in roots:
        if not os.path.isdir(r):
            continue
        for dirpath, dirnames, filenames in os.walk(r):
            if time.time() > deadline:
                timed_out = True
                break
            dirnames[:] = [d for d in dirnames
                           if d.lower() not in _SKIP_DIRS and not d.startswith(".")]
            for fn in filenames:
                if fnmatch.fnmatch(fn.lower(), pattern):
                    p = os.path.join(dirpath, fn)
                    try:
                        kb = os.path.getsize(p) // 1024
                    except OSError:
                        kb = 0
                    matches.append(f"{p}  ({kb} KB)")
                    if len(matches) >= max_results:
                        return ("Found (showing first "
                                f"{max_results} — refine to narrow down):\n"
                                + "\n".join(matches))
        if timed_out:
            break
    if not matches:
        hint = " (search timed out — try a folder: [FIND: name | C:\\path])" if timed_out else ""
        return f"No files matching '{query}' found in your user folders{hint}."
    note = "\n[search timed out — more may exist]" if timed_out else ""
    return "Found:\n" + "\n".join(matches) + note


def get_file_tool_description():
    return ("\n[FILE: /path/to/file]   → Read file (PDF/docx/xlsx/txt/code)"
            "\n[WRITE: path | content] → Create/save a file (.docx, .md, .txt, code…). "
            "Relative names go to Desktop. Put the FULL file content after the |."
            "\n[FIND: name or *.pdf]   → Find files by name in the user's folders. "
            "Optionally [FIND: pattern | C:\\folder]\n")
